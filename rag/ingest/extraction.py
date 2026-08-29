"""Format-specific extraction: raw file -> text / transcript / audio windows.

ARCHITECTURE.md §2:
  DOCX  -> python-docx
  PDF   -> PyMuPDF text layer; pytesseract OCR fallback for scanned pages
  Image -> pytesseract OCR (stored as text_preview only; ImageBind embeds the
           image itself), so text-heavy screenshots match reliably
  Audio -> Whisper (small) transcript + segment timestamps for citation
           display; sliced into fixed ~2s windows for ImageBind embedding

Everything here returns plain dataclasses; embedding and Chroma writes happen
in rag/ingest/corpus.py.
"""

from __future__ import annotations

import logging
import subprocess
import tempfile
import wave
from dataclasses import dataclass, field
from pathlib import Path

from rag import config

logger = logging.getLogger(__name__)

DOC_EXTS = {".docx", ".doc"}
PDF_EXTS = {".pdf"}
TEXT_EXTS = {".txt", ".md"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp", ".gif"}
AUDIO_EXTS = {".wav", ".mp3", ".m4a", ".flac", ".ogg", ".aac", ".wma"}

# Below this many characters, a PDF page's text layer is treated as empty
# (scanned page) and OCR is attempted instead.
PDF_OCR_CHAR_THRESHOLD = 40


@dataclass
class PageText:
    text: str
    page_number: int | None = None


@dataclass
class DocumentExtraction:
    modality: str          # "document"
    pages: list[PageText]
    used_ocr: bool = False


@dataclass
class ImageExtraction:
    modality: str          # "image"
    ocr_text: str
    width: int | None = None
    height: int | None = None


@dataclass
class AudioSegment:
    text: str
    start_time: float
    end_time: float


@dataclass
class AudioWindow:
    path: Path             # temp wav, one ~2s clip
    start_time: float
    end_time: float


@dataclass
class AudioExtraction:
    modality: str          # "audio"
    segments: list[AudioSegment]        # Whisper transcript segments (citation)
    windows: list[AudioWindow]          # fixed windows for ImageBind embedding
    full_text: str
    duration: float
    _tmpdir: tempfile.TemporaryDirectory | None = field(default=None, repr=False)

    def cleanup(self) -> None:
        if self._tmpdir is not None:
            self._tmpdir.cleanup()
            self._tmpdir = None


# --------------------------------------------------------------------------
# OCR
# --------------------------------------------------------------------------

_ocr_checked = False
_ocr_available = False


def ocr_available() -> bool:
    global _ocr_checked, _ocr_available
    if _ocr_checked:
        return _ocr_available
    _ocr_checked = True
    try:
        import pytesseract
        from PIL import Image  # noqa: F401

        if config.TESSERACT_CMD:
            pytesseract.pytesseract.tesseract_cmd = config.TESSERACT_CMD
        pytesseract.get_tesseract_version()
        _ocr_available = True
    except Exception as e:  # noqa: BLE001
        logger.warning(
            "Tesseract OCR unavailable (%s). Scanned PDFs and image text "
            "previews will be empty. Install the Tesseract binary + `pip "
            "install pytesseract pillow`.",
            e,
        )
        _ocr_available = False
    return _ocr_available


def _ocr_image(pil_image) -> str:
    import pytesseract

    try:
        return pytesseract.image_to_string(pil_image).strip()
    except Exception as e:  # noqa: BLE001
        logger.warning("OCR failed on an image: %s", e)
        return ""


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------

def extract_docx(path: Path) -> DocumentExtraction:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError("python-docx not installed (`pip install python-docx`)") from e

    if path.suffix.lower() == ".doc":
        raise RuntimeError(
            f"{path.name}: legacy .doc not supported directly — convert to .docx "
            "(e.g. libreoffice --headless --convert-to docx)."
        )

    document = docx.Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full = "\n".join(parts)
    # python-docx exposes no page breaks reliably; treat a DOCX as one page.
    return DocumentExtraction(modality="document", pages=[PageText(full, page_number=None)])


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------

def extract_text_file(path: Path) -> DocumentExtraction:
    text = path.read_text(encoding="utf-8", errors="replace")
    return DocumentExtraction(
        modality="document", pages=[PageText(text=text, page_number=None)]
    )


def extract_pdf(path: Path) -> DocumentExtraction:
    try:
        import pymupdf as fitz
    except ImportError as e:
        raise RuntimeError("PyMuPDF not installed (`pip install pymupdf`)") from e

    doc = fitz.open(str(path))
    pages: list[PageText] = []
    used_ocr = False
    try:
        for i, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if len(text) < PDF_OCR_CHAR_THRESHOLD and ocr_available():
                text = _ocr_pdf_page(page) or text
                if text:
                    used_ocr = True
            if text.strip():
                pages.append(PageText(text=text, page_number=i))
    finally:
        doc.close()

    if not pages:
        logger.warning("%s: no extractable text (even after OCR attempt)", path.name)
    return DocumentExtraction(modality="document", pages=pages, used_ocr=used_ocr)


def _ocr_pdf_page(page) -> str:
    import io

    from PIL import Image

    pix = page.get_pixmap(dpi=200)
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    return _ocr_image(img)


# --------------------------------------------------------------------------
# Image
# --------------------------------------------------------------------------

def extract_image(path: Path) -> ImageExtraction:
    width = height = None
    ocr_text = ""
    try:
        from PIL import Image

        with Image.open(str(path)) as img:
            width, height = img.size
            if ocr_available():
                ocr_text = _ocr_image(img.convert("RGB"))
    except Exception as e:  # noqa: BLE001
        logger.warning("Could not open image %s: %s", path.name, e)
    return ImageExtraction(
        modality="image", ocr_text=ocr_text, width=width, height=height
    )


# --------------------------------------------------------------------------
# Audio
# --------------------------------------------------------------------------

_whisper_model = None


def _load_whisper():
    global _whisper_model
    if _whisper_model is not None:
        return _whisper_model
    try:
        import whisper  # openai-whisper
    except ImportError as e:
        raise RuntimeError(
            "openai-whisper not installed (`pip install openai-whisper`)"
        ) from e
    from rag.core.embeddings import resolve_device

    device = resolve_device()
    config.WHISPER_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    logger.info("Loading Whisper (%s) on %s ...", config.WHISPER_MODEL, device)
    _whisper_model = whisper.load_model(
        config.WHISPER_MODEL, device=device, download_root=str(config.WHISPER_CACHE_DIR)
    )
    return _whisper_model


def unload_whisper() -> None:
    global _whisper_model
    _whisper_model = None


def _load_audio_array(path: Path):
    """Decode any supported audio file to mono float32 at AUDIO_SAMPLE_RATE."""
    import numpy as np

    try:
        import soundfile as sf
        import librosa

        data, sr = sf.read(str(path), dtype="float32", always_2d=False)
        if data.ndim > 1:
            data = data.mean(axis=1)
        if sr != config.AUDIO_SAMPLE_RATE:
            data = librosa.resample(
                data, orig_sr=sr, target_sr=config.AUDIO_SAMPLE_RATE
            )
        return data.astype("float32")
    except Exception as e:  # noqa: BLE001 - fall back to whisper's ffmpeg loader
        logger.debug("soundfile/librosa load failed (%s); using whisper.load_audio", e)
        import whisper

        return whisper.load_audio(str(path), sr=config.AUDIO_SAMPLE_RATE)


def _write_wav(samples, sr: int, out_path: Path) -> None:
    import numpy as np

    clipped = np.clip(samples, -1.0, 1.0)
    pcm = (clipped * 32767).astype("<i2")
    with wave.open(str(out_path), "wb") as w:
        w.setnchannels(1)
        w.setsampwidth(2)
        w.setframerate(sr)
        w.writeframes(pcm.tobytes())


def slice_audio_windows(
    path: Path, tmpdir: Path, samples=None
) -> tuple[list[AudioWindow], float]:
    """Cut a recording into consecutive fixed windows (ARCHITECTURE.md §2).

    Pass `samples` (float32 mono @ AUDIO_SAMPLE_RATE) to skip re-decoding.
    """
    if samples is None:
        samples = _load_audio_array(path)
    sr = config.AUDIO_SAMPLE_RATE
    duration = len(samples) / sr
    win = int(config.AUDIO_WINDOW_SECONDS * sr)
    hop = int(config.AUDIO_WINDOW_HOP_SECONDS * sr)

    windows: list[AudioWindow] = []
    idx = 0
    start = 0
    while start < len(samples):
        seg = samples[start : start + win]
        if len(seg) < int(0.25 * sr):  # skip a sub-250ms trailing sliver
            break
        if len(seg) < win:  # pad the final short window to a full clip
            import numpy as np

            seg = np.pad(seg, (0, win - len(seg)))
        out = tmpdir / f"win_{idx:05d}.wav"
        _write_wav(seg, sr, out)
        windows.append(
            AudioWindow(
                path=out,
                start_time=round(start / sr, 3),
                end_time=round(min(start / sr + config.AUDIO_WINDOW_SECONDS, duration), 3),
            )
        )
        idx += 1
        start += hop
    return windows, duration


def _sidecar_segments(path: Path) -> list[AudioSegment] | None:
    """A `<audio>.segments.json` next to the file supplies a ready-made,
    time-aligned transcript (list of {text, start, end}) — used for corpus
    audio that already ships with a real transcription (e.g. the AMI Meeting
    Corpus), so ingestion skips Whisper entirely. Returns None if absent."""
    sc = path.with_name(path.name + ".segments.json")
    if not sc.exists():
        return None
    try:
        import json
        data = json.loads(sc.read_text(encoding="utf-8"))
        segs = [
            AudioSegment(
                text=str(d.get("text", "")).strip(),
                start_time=round(float(d["start"]), 3),
                end_time=round(float(d["end"]), 3),
            )
            for d in data
            if str(d.get("text", "")).strip()
        ]
        logger.info("Using sidecar transcript for %s (%d segments, no Whisper)",
                    path.name, len(segs))
        return segs
    except Exception as e:  # noqa: BLE001
        logger.warning("sidecar %s unreadable (%s) — falling back to Whisper", sc.name, e)
        return None


def extract_audio(path: Path) -> AudioExtraction:
    segments = _sidecar_segments(path)
    if segments is not None:
        samples = _load_audio_array(path)
        full_text = " ".join(s.text for s in segments)
    else:
        model = _load_whisper()
        logger.info("Transcribing %s ...", path.name)
        # Decode ourselves (soundfile/librosa) rather than letting Whisper shell
        # out to an ffmpeg binary — ffmpeg may not be on PATH, and for WAV it's
        # not needed. _load_audio_array returns float32 mono @ 16 kHz, which is
        # exactly what whisper.transcribe expects as an ndarray.
        samples = _load_audio_array(path)
        result = model.transcribe(samples, fp16=(_whisper_uses_fp16()))
        segments = [
            AudioSegment(
                text=s["text"].strip(),
                start_time=round(float(s["start"]), 3),
                end_time=round(float(s["end"]), 3),
            )
            for s in result.get("segments", [])
            if s.get("text", "").strip()
        ]
        full_text = result.get("text", "").strip()

    # Only cut 2 s windows if something will embed them (config.AUDIO_EMBED_WINDOWS
    # on the ingest side, or a query-time audio clip). Skipping this saves the
    # per-window WAV writes on every audio file.
    duration = len(samples) / config.AUDIO_SAMPLE_RATE
    windows: list[AudioWindow] = []
    tmp = None
    if config.AUDIO_EMBED_WINDOWS:
        tmp = tempfile.TemporaryDirectory(prefix="rag_audio_")
        windows, duration = slice_audio_windows(path, Path(tmp.name), samples=samples)

    return AudioExtraction(
        modality="audio",
        segments=segments,
        windows=windows,
        full_text=full_text,
        duration=duration,
        _tmpdir=tmp,
    )


def _whisper_uses_fp16() -> bool:
    from rag.core.embeddings import resolve_device

    return resolve_device() == "cuda"


# --------------------------------------------------------------------------
# Dispatch
# --------------------------------------------------------------------------

def classify(path: Path) -> str | None:
    ext = path.suffix.lower()
    if ext in DOC_EXTS or ext in PDF_EXTS or ext in TEXT_EXTS:
        return "document"
    if ext in IMAGE_EXTS:
        return "image"
    if ext in AUDIO_EXTS:
        return "audio"
    return None


def segment_text_for_time(segments: list[AudioSegment], start: float, end: float) -> str:
    """Transcript text overlapping a [start, end) embedding window — used as the
    text_preview for audio chunks so citations show readable words."""
    hits = [s.text for s in segments if s.end_time > start and s.start_time < end]
    return " ".join(hits).strip()
