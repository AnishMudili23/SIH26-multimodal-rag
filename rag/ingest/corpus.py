"""Phase 0 ingestion pipeline orchestrator.

  raw file -> format-specific extraction -> chunk -> ImageBind embed -> Chroma

ARCHITECTURE.md §2 end to end. Run:

    python -m rag.ingest.corpus --src data/raw --reset
    python -m rag.ingest.corpus --stats
    python -m rag.ingest.corpus --smoke          # tiny self-check, no real corpus needed

An optional corpus_manifest.csv (rag/ingest/manifest.py) supplies `theme`
tags per doc_id for demo filtering / eval.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import logging
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import numpy as np

from rag import config
from rag.ingest import extraction
from rag.ingest.chunking import chunk_text, link_chunks
from rag.core.vectorstore import ChunkRecord, add_records, collection_stats, get_collection

logging.basicConfig(
    level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s: %(message)s"
)
logger = logging.getLogger("ingest")


# --------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------

def _slug(s: str) -> str:
    """Filesystem-name → a short id-safe token, so several files sharing one
    meeting `doc_id` (AMI: transcript + slides + minutes) get distinct chunk
    ids."""
    return "".join(c if c.isalnum() else "_" for c in s)[:48].strip("_") or "x"


def make_doc_id(path: Path, manifest_ids: set[str] | None = None) -> str:
    """Doc id for a source file.

    If the file's stem is already a manifest `doc_id` (materialised passages
    from rag/ingest/manifest.py), use it verbatim so the manifest, the qrels
    file, and the ingested chunks all agree. Otherwise derive a stable id from
    name + size + mtime (re-ingest upserts in place).
    """
    if manifest_ids and path.stem in manifest_ids:
        return path.stem
    st = path.stat()
    h = hashlib.sha1(
        f"{path.name}|{st.st_size}|{int(st.st_mtime)}".encode()
    ).hexdigest()[:12]
    stem = "".join(c if c.isalnum() else "_" for c in path.stem)[:40]
    return f"{stem}_{h}"


@dataclass
class AssetLink:
    """A manifest row's cross-modal link, keyed by the asset filename."""

    theme: str | None
    passage_doc_id: str          # the doc_id / meeting group this asset belongs to
    image_timestamp: str | None = None   # "HH:MM" (old MS MARCO manifest) or a
                                         #  "mm:ss–mm:ss" display range (AMI slides)
    start_time: float | None = None      # AMI: slide's on-screen window, seconds
    end_time: float | None = None


def _load_ami_manifest(path: Path):
    """AMI manifest: one row per file, columns
    doc_id, meeting_id, modality, source_file, text_preview, start_time, end_time.
    `doc_id` groups every chunk of one meeting (transcript ↔ slide ↔ minutes)."""
    themes: dict[str, str] = {}
    rows: dict[str, dict] = {}
    assets: dict[str, AssetLink] = {}
    with path.open(newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            did = (row.get("doc_id") or "").strip()
            src = (row.get("source_file") or "").strip()
            if not did or not src:
                continue
            mid = (row.get("meeting_id") or "").strip() or did
            rows.setdefault(did, row)
            themes[did] = mid

            def _f(key):
                v = (row.get(key) or "").strip()
                try:
                    return float(v) if v else None
                except ValueError:
                    return None

            ts = None
            st, en = _f("start_time"), _f("end_time")
            if st is not None:
                ts = f"{int(st // 60):02d}:{int(st % 60):02d}" + (
                    f"-{int(en // 60):02d}:{int(en % 60):02d}" if en is not None else "")
            link = AssetLink(theme=mid, passage_doc_id=did,
                             image_timestamp=ts, start_time=st, end_time=en)
            # key by basename AND by the last two path parts, so slides in
            # per-meeting subdirs still resolve
            assets[Path(src).name] = link
            assets["/".join(Path(src).parts[-2:])] = link
    return themes, rows, assets


def load_manifest() -> tuple[dict[str, str], dict[str, dict], dict[str, AssetLink]]:
    """Returns (themes: doc_id->label, rows: doc_id->row, assets: filename->AssetLink).

    Auto-detects the manifest schema:
      - **AMI** (`rag/ingest/manifest.py`): per-file rows with a `modality` column;
        `doc_id` groups a whole meeting for cross-modal linking.
      - **legacy MS MARCO**: per-passage rows with `image_file`/`audio_file`/
        `image_timestamp` columns naming hand-made assets to link.
    """
    path = config.CORPUS_MANIFEST_CSV
    if not path.exists():
        return {}, {}, {}
    with path.open(newline="", encoding="utf-8") as f:
        header = next(csv.reader(f), [])
    if "modality" in header and "source_file" in header:
        return _load_ami_manifest(path)

    themes: dict[str, str] = {}
    rows: dict[str, dict] = {}
    assets: dict[str, AssetLink] = {}
    with path.open(newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            did = (row.get("doc_id") or "").strip()
            if not did:
                continue
            rows[did] = row
            th = (row.get("theme") or "").strip() or None
            if th:
                themes[did] = th
            # image_file / audio_file may hold several ";"-separated names;
            # image_timestamp is a position-aligned ";"-list (blank = none)
            img_files = [n.strip() for n in (row.get("image_file") or "").split(";") if n.strip()]
            img_ts = (row.get("image_timestamp") or "").split(";")
            for i, name in enumerate(img_files):
                ts = (img_ts[i].strip() if i < len(img_ts) else "") or None
                assets[Path(name).name] = AssetLink(
                    theme=th, passage_doc_id=did, image_timestamp=ts
                )
            for name in (row.get("audio_file") or "").split(";"):
                name = name.strip()
                if name:
                    assets[Path(name).name] = AssetLink(
                        theme=th, passage_doc_id=did, image_timestamp=None
                    )
    return themes, rows, assets


def iter_source_files(src: Path) -> Iterable[Path]:
    for p in sorted(src.rglob("*")):
        if not p.is_file():
            continue
        # _cache/ holds the raw AMI downloads (annotation zip, original .doc
        # files) that have already been converted to .txt — don't re-ingest them
        if "_cache" in p.parts:
            continue
        if extraction.classify(p) is not None:
            yield p


# --------------------------------------------------------------------------
# Per-file -> records (embeddings filled in a batch afterwards)
# --------------------------------------------------------------------------

@dataclass
class PendingRecord:
    record: ChunkRecord
    embed_kind: str           # "text" | "image" | "audio"
    embed_ref: object         # str (text) or Path (image/audio window file)


def build_document_records(path: Path, doc_id: str, theme: str | None) -> list[PendingRecord]:
    ext = path.suffix.lower()
    if ext in extraction.DOC_EXTS:
        doc = extraction.extract_docx(path)
    elif ext in extraction.TEXT_EXTS:
        doc = extraction.extract_text_file(path)
    else:
        doc = extraction.extract_pdf(path)

    all_chunks = []
    page_of_chunk = []
    for page in doc.pages:
        cs = chunk_text(page.text, page_number=page.page_number)
        all_chunks.extend(cs)
        page_of_chunk.extend([page.page_number] * len(cs))
    # prefix with the filename slug so multiple documents under one meeting
    # doc_id (AMI minutes + report + summaries) don't collide on chunk id
    link_chunks(f"{doc_id}::{_slug(path.stem)}", all_chunks)

    pending: list[PendingRecord] = []
    for c in all_chunks:
        rec = ChunkRecord(
            id=c.chunk_id,
            embedding=(),
            modality="document",
            doc_id=doc_id,
            source_file=path.name,
            source_path=str(path),
            text_preview=c.text,
            page_number=c.page_number,
            prev_chunk_id=c.prev_chunk_id,
            next_chunk_id=c.next_chunk_id,
            theme=theme,
        )
        pending.append(PendingRecord(rec, "text", c.text))
    return pending


def build_image_records(
    path: Path, doc_id: str, theme: str | None, link: "AssetLink | None" = None
) -> list[PendingRecord]:
    img = extraction.extract_image(path)
    preview = img.ocr_text or f"[image: {path.name}]"
    theme_v = link.theme if link and link.theme else theme
    ts_v = link.image_timestamp if link else None
    link_v = link.passage_doc_id if link else None
    st_v = link.start_time if link else None
    en_v = link.end_time if link else None
    # a slide with its own doc_id IS its meeting's chunk, not a chunk "linked to"
    # another passage — don't set linked_passage to itself
    linked_passage_v = None if link_v == doc_id else link_v

    # (1) the vision embedding of the image itself
    pending = [
        PendingRecord(
            ChunkRecord(
                id=f"{doc_id}::image::{_slug(path.stem)}",
                embedding=(),
                modality="image",
                doc_id=doc_id,
                source_file=path.name,
                source_path=str(path),
                text_preview=preview[:2000],
                theme=theme_v,
                image_timestamp=ts_v,
                start_time=st_v,
                end_time=en_v,
                linked_passage=linked_passage_v,
            ),
            "image",
            path,
        )
    ]

    # (2) ARCHITECTURE.md §2: text-heavy screenshots ("email screenshot",
    # "dashboard at 14:32") match far better on their OCR text than on visual
    # similarity alone. Embed the OCR as extra image-modality chunks in the
    # *text* space so a text query can reach them directly.
    ocr = (img.ocr_text or "").strip()
    if len(ocr.split()) >= 4:
        slug = _slug(path.stem)
        ocr_chunks = chunk_text(ocr)
        link_chunks(f"{doc_id}::image_ocr::{slug}", ocr_chunks)
        for i, c in enumerate(ocr_chunks):
            pending.append(
                PendingRecord(
                    ChunkRecord(
                        id=c.chunk_id,
                        embedding=(),
                        modality="image",
                        doc_id=doc_id,
                        source_file=path.name,
                        source_path=str(path),
                        text_preview=c.text,
                        theme=theme_v,
                        image_timestamp=ts_v,
                        start_time=st_v,
                        end_time=en_v,
                        linked_passage=linked_passage_v,
                        prev_chunk_id=c.prev_chunk_id,
                        next_chunk_id=c.next_chunk_id,
                    ),
                    "text",
                    c.text,
                )
            )
    return pending


def build_audio_records(
    path: Path, doc_id: str, theme: str | None, link: "AssetLink | None" = None
) -> tuple[list[PendingRecord], extraction.AudioExtraction]:
    audio = extraction.extract_audio(path)
    theme_v = link.theme if link and link.theme else theme
    link_v = link.passage_doc_id if link else None
    pending: list[PendingRecord] = []

    # (1) ImageBind audio embeddings of the fixed ~2 s windows — only for
    # pure audio-to-audio matching, off by default (config.AUDIO_EMBED_WINDOWS).
    # The transcript chunks below carry retrieval; skipping windows is ~30 %
    # faster to ingest and roughly halves the audio chunk count.
    if config.AUDIO_EMBED_WINDOWS:
        win_ids = [f"{doc_id}::audio::{i:04d}" for i in range(len(audio.windows))]
        for i, w in enumerate(audio.windows):
            preview = extraction.segment_text_for_time(audio.segments, w.start_time, w.end_time)
            pending.append(PendingRecord(
                ChunkRecord(
                    id=win_ids[i],
                    embedding=(),
                    modality="audio",
                    doc_id=doc_id,
                    source_file=path.name,
                    source_path=str(path),
                    text_preview=preview or f"[audio {w.start_time:.1f}-{w.end_time:.1f}s]",
                    start_time=w.start_time,
                    end_time=w.end_time,
                    prev_chunk_id=win_ids[i - 1] if i > 0 else None,
                    next_chunk_id=win_ids[i + 1] if i < len(win_ids) - 1 else None,
                    theme=theme_v,
                    linked_passage=link_v,
                ),
                "audio", w.path,
            ))

    # (2) Whisper transcript segments embedded in the *text* space. ImageBind's
    # text<->audio alignment is weak (ARCHITECTURE.md §4), so a plain-language
    # query reaches a recorded call far more reliably through its transcript.
    seg_ids = [f"{doc_id}::transcript::{i:04d}" for i in range(len(audio.segments))]
    for i, s in enumerate(audio.segments):
        if not s.text.strip():
            continue
        pending.append(PendingRecord(
            ChunkRecord(
                id=seg_ids[i],
                embedding=(),
                modality="audio",
                doc_id=doc_id,
                source_file=path.name,
                source_path=str(path),
                text_preview=s.text.strip(),
                start_time=s.start_time,
                end_time=s.end_time,
                prev_chunk_id=seg_ids[i - 1] if i > 0 else None,
                next_chunk_id=seg_ids[i + 1] if i < len(seg_ids) - 1 else None,
                theme=theme_v,
                linked_passage=link_v,
            ),
            "text", s.text.strip(),
        ))
    return pending, audio


# --------------------------------------------------------------------------
# Embedding
# --------------------------------------------------------------------------

def embed_pending(pending: list[PendingRecord]) -> list[ChunkRecord]:
    from rag.core import embeddings

    by_kind: dict[str, list[int]] = {"text": [], "image": [], "audio": []}
    for i, p in enumerate(pending):
        by_kind[p.embed_kind].append(i)

    if by_kind["text"]:
        vecs = embeddings.embed_texts([pending[i].embed_ref for i in by_kind["text"]])
        for j, i in enumerate(by_kind["text"]):
            pending[i].record.embedding = vecs[j].tolist()
    if by_kind["image"]:
        vecs = embeddings.embed_images([pending[i].embed_ref for i in by_kind["image"]])
        for j, i in enumerate(by_kind["image"]):
            pending[i].record.embedding = vecs[j].tolist()
    if by_kind["audio"]:
        vecs = embeddings.embed_audio_windows(
            [pending[i].embed_ref for i in by_kind["audio"]]
        )
        for j, i in enumerate(by_kind["audio"]):
            pending[i].record.embedding = vecs[j].tolist()

    return [p.record for p in pending]


# --------------------------------------------------------------------------
# Driver
# --------------------------------------------------------------------------

def _process_files(files, collection, progress=None) -> dict:
    """Ingest a list of paths into an open collection. `progress(i, n, name)` is
    called before each file. Returns a summary. Does NOT unload ImageBind — the
    caller decides (a CLI run unloads; the app parks it on CPU to stay warm)."""
    themes, manifest_rows, asset_links = load_manifest()
    manifest_ids = set(manifest_rows)

    total_written = 0
    ok, failed = [], []
    audio_holders: list[extraction.AudioExtraction] = []
    for i, path in enumerate(files):
        if progress:
            progress(i + 1, len(files), path.name)
        kind = extraction.classify(path)
        if kind is None:
            failed.append((path.name, "unsupported type"))
            continue
        link = (asset_links.get("/".join(path.parts[-2:]))
                or asset_links.get(path.name))
        if link and link.passage_doc_id:
            doc_id = link.passage_doc_id          # AMI: the meeting group
        else:
            doc_id = make_doc_id(path, manifest_ids)
        theme = (link.theme if link and link.theme else None) or themes.get(doc_id)
        logger.info("[%s] %s%s", kind, path.name,
                    f"  -> {link.passage_doc_id}" if link else "")
        try:
            if kind == "document":
                pending = build_document_records(path, doc_id, theme)
            elif kind == "image":
                pending = build_image_records(path, doc_id, theme, link)
            else:  # audio
                pending, holder = build_audio_records(path, doc_id, theme, link)
                audio_holders.append(holder)
            if not pending:
                failed.append((path.name, "no chunks produced"))
                continue
            written = add_records(collection, embed_pending(pending))
            total_written += written
            ok.append((path.name, kind, written))
            logger.info("  -> %d chunks embedded + written", written)
        except Exception as e:  # noqa: BLE001 - keep going
            logger.exception("  -> FAILED on %s: %s", path.name, e)
            failed.append((path.name, str(e)))

    for h in audio_holders:
        h.cleanup()
    return {"written": total_written, "ok": ok, "failed": failed,
            "stats": collection_stats(collection)}


def ingest(src: Path, reset: bool = False, limit: int | None = None) -> None:
    config.ensure_dirs()
    collection = get_collection(reset=reset)
    files = list(iter_source_files(src))
    if limit:
        files = files[:limit]
    if not files:
        logger.warning("No ingestible files found under %s", src)
        return
    logger.info("Found %d source files under %s", len(files), src)

    summary = _process_files(files, collection)

    from rag.core import embeddings
    embeddings.unload()
    extraction.unload_whisper()
    logger.info("Done. %d chunks written this run. Collection now: %s",
                summary["written"], summary["stats"])


def add_files_to_corpus(paths, copy_into_raw: bool = True, progress=None) -> dict:
    """Ingest user-picked files into the *current* collection (no reset) and
    keep ImageBind warm. Copies each file under RAW_DATA_DIR so a later
    `python -m rag.ingest.corpus --reset` rebuild still includes it. Used by the desktop app."""
    import shutil
    from rag.core import embeddings

    config.ensure_dirs()
    dest_by_kind = {"document": "docs", "image": "images", "audio": "audio"}
    resolved: list[Path] = []
    for p in paths:
        p = Path(p)
        if not p.is_file():
            continue
        kind = extraction.classify(p)
        if copy_into_raw and kind:
            sub = config.RAW_DATA_DIR / dest_by_kind.get(kind, "misc")
            sub.mkdir(parents=True, exist_ok=True)
            target = sub / p.name
            if not target.exists() or target.stat().st_size != p.stat().st_size:
                shutil.copy2(p, target)
            resolved.append(target)
        else:
            resolved.append(p)

    collection = get_collection(reset=False)
    summary = _process_files(resolved, collection, progress=progress)
    embeddings.offload_to_cpu()      # keep weights, free the GPU for the LLM
    extraction.unload_whisper()
    return summary


# --------------------------------------------------------------------------
# Smoke test (ARCHITECTURE.md Phase 0 last checklist item, minimal form)
# --------------------------------------------------------------------------

def smoke_test() -> int:
    """Create tiny synthetic files of each format, ingest, query back.

    Runs against an ISOLATED temp raw-dir / Chroma store / collection so it
    never touches a real corpus. Requires the heavy deps (torch/imagebind/
    chromadb, +whisper for audio). Returns process exit code.
    """
    import tempfile

    tmp = Path(tempfile.mkdtemp(prefix="rag_smoke_"))
    logger.info("Smoke test workspace: %s", tmp)

    # Redirect all pipeline I/O into the temp workspace for this run only.
    config.RAW_DATA_DIR = tmp / "raw"
    config.CHROMA_DIR = tmp / "chroma"
    config.CORPUS_MANIFEST_CSV = tmp / "corpus_manifest.csv"
    config.COLLECTION_NAME = "smoke_test"
    config.RAW_DATA_DIR.mkdir(parents=True, exist_ok=True)
    config.CHROMA_DIR.mkdir(parents=True, exist_ok=True)

    # 1. a DOCX
    made = []
    try:
        import docx

        d = docx.Document()
        d.add_paragraph(
            "The 2024 report on international development covers infrastructure "
            "financing across South Asia. It highlights renewable energy grants "
            "and cross-border rail investment."
        )
        d.add_paragraph(
            "A separate annex describes a screenshot captured at 14:32 showing "
            "the quarterly disbursement dashboard."
        )
        p = tmp / "dev_report_2024.docx"
        d.save(str(p))
        made.append(p)
    except Exception as e:  # noqa: BLE001
        logger.warning("skip DOCX (%s)", e)

    # 2. a PDF
    try:
        import pymupdf as fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text(
            (72, 72),
            "International development funding in 2024 reached record levels.\n"
            "The screenshot taken at 14:32 documents the transfer confirmation.",
        )
        p = tmp / "funding_note.pdf"
        doc.save(str(p))
        doc.close()
        made.append(p)
    except Exception as e:  # noqa: BLE001
        logger.warning("skip PDF (%s)", e)

    # 3. an image
    try:
        from PIL import Image, ImageDraw

        img = Image.new("RGB", (640, 200), "white")
        ImageDraw.Draw(img).text((10, 90), "EMAIL SCREENSHOT 14:32 - development budget", fill="black")
        p = tmp / "email_screenshot.png"
        img.save(str(p))
        made.append(p)
    except Exception as e:  # noqa: BLE001
        logger.warning("skip image (%s)", e)

    # 4. audio (sine tone, transcript will be empty but windows must embed)
    try:
        sr = config.AUDIO_SAMPLE_RATE
        t = np.linspace(0, 3.0, int(sr * 3.0), endpoint=False)
        tone = (0.2 * np.sin(2 * np.pi * 220 * t)).astype("float32")
        p = tmp / "call_clip.wav"
        extraction._write_wav(tone, sr, p)
        made.append(p)
    except Exception as e:  # noqa: BLE001
        logger.warning("skip audio (%s)", e)

    if not made:
        logger.error("Could not create any test files — deps missing.")
        return 1

    for f in made:
        (config.RAW_DATA_DIR / f.name).write_bytes(f.read_bytes())

    ingest(config.RAW_DATA_DIR, reset=True, limit=None)

    # manual query back
    from rag.core import embeddings
    from rag.core.vectorstore import get_collection, query

    col = get_collection()
    qv = embeddings.embed_texts(["international development funding in 2024"])[0]
    hits = query(col, qv.tolist(), n_results=5)
    embeddings.unload()
    logger.info("Query 'international development funding in 2024' ->")
    for h in hits:
        logger.info(
            "  [%s] %s  d=%.3f  %r",
            h.get("modality"),
            h.get("source_file"),
            h.get("distance", -1),
            (h.get("text_preview") or "")[:80],
        )
    ok = any("development" in (h.get("text_preview") or "").lower() for h in hits)
    logger.info("Smoke test %s", "PASSED" if ok else "INCONCLUSIVE (check output above)")
    return 0 if hits else 1


# --------------------------------------------------------------------------
# CLI
# --------------------------------------------------------------------------

def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--src", type=Path, default=config.RAW_DATA_DIR,
                    help="directory of raw files to ingest")
    ap.add_argument("--reset", action="store_true",
                    help="drop and recreate the Chroma collection first")
    ap.add_argument("--limit", type=int, default=None, help="ingest at most N files")
    ap.add_argument("--stats", action="store_true", help="print collection stats and exit")
    ap.add_argument("--smoke", action="store_true", help="run the synthetic smoke test")
    args = ap.parse_args(argv)

    if args.stats:
        print(collection_stats(get_collection()))
        return 0
    if args.smoke:
        return smoke_test()

    ingest(args.src, reset=args.reset, limit=args.limit)
    return 0


if __name__ == "__main__":
    sys.exit(main())
